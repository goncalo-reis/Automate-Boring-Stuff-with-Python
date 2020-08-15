import docx
from docx.enum.text import WD_BREAK

doc = docx.Document('atum.docx')
with open('guests.txt') as file:
    names = file.read().split('\n')

for name in names:
    doc.add_paragraph('It would be a pleasure to have the company of', style='Signature')
    doc.add_paragraph(name, style='Title')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    doc.add_paragraph('April 1st')
    date_line = doc.add_paragraph('at 7 o\'clock')
    date_line.runs[0].add_break(break_type=WD_BREAK.PAGE)

doc.save('invitations.docx')
